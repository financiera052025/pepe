import streamlit as st
import re
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuración de la página web
st.set_page_config(page_title="Analizador Financiero", layout="centered")
st.title("INSTRUMENTOS FINANCIEROS")
st.write("Selecciona el módulo de procesamiento que deseas aplicar.")

# Diccionario Global
DICCIONARIO_BANCOS = {
    "0102": "BANCO DE VENEZUELA", "0104": "BANCO VENEZOLANO DE CREDITO",
    "0105": "BANCO MERCANTIL", "0108": "BANCO PROVINCIAL, S.A.",
    "0114": "BANCO DEL CARIBE", "0115": "BANCO EXTERIOR",
    "0128": "BANCO CARONI", "0134": "BANESCO", "0151": "FONDO COMUN",
    "0156": "100% BANCO", "0157": "DELSUR", "0163": "BANCO DEL TESORO",
    "0166": "BANCO AGRICOLA DE VENEZUELA", "0168": "BANCRECER",
    "0169": "MI BANCO", "0172": "BANCAMIGA", "0174": "BANPLUS",
    "0175": "BANCO DIGITAL DE LOS TRABAJADORES", "0177": "BANCO BANFANB",
    "0191": "BANCO NACIONAL DE CREDITO (BNC)"
}

# --- FUNCIONES DE AYUDA GLOBALES ---
def obtener_banco_por_cuenta(nro_cuenta):
    cuenta_limpia = re.sub(r'\D', '', str(nro_cuenta))
    codigo = cuenta_limpia[:4]
    return DICCIONARIO_BANCOS.get(codigo, "OTRA ENTIDAD BANCARIA")

def limpiar_monto_venezuela(monto_str):
    if not monto_str:
        return 0.0
    m_limpio = str(monto_str).replace('"', '').replace(' ', '').strip()
    
    # Lógica avanzada del módulo 3 para mayor compatibilidad
    if ',' in m_limpio and '.' in m_limpio:
        if m_limpio.rfind(',') > m_limpio.rfind('.'): 
            m_limpio = m_limpio.replace('.', '').replace(',', '.')
        else: 
            m_limpio = m_limpio.replace(',', '')
    elif ',' in m_limpio and '.' not in m_limpio:
        m_limpio = m_limpio.replace(',', '.')
    
    m_limpio = re.sub(r'[^\d\.]', '', m_limpio)
    try:
        return float(m_limpio)
    except ValueError:
        return 0.0

def formatear_monto_venezuela(monto_float):
    return f"{monto_float:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

def formatear_titulo_persona(nombre, identificador):
    nombre = nombre.strip().upper()
    identificador = identificador.strip().upper()
    if identificador.startswith(('J', 'G')) or "DISTRIBUIDORA" in nombre or "C.A" in nombre:
        return f"INSTRUMENTOS FINANCIEROS DE LA SOCIEDAD MERCANTIL {nombre}, RIF {identificador}:"
    primer_nombre = nombre.split()[0] if nombre.split() else ""
    if primer_nombre.endswith(('A', 'ITH', 'Y', 'IS', 'NID')) and not primer_nombre.endswith(('JUAN', 'ANGEL', 'JOSUE')):
        return f"INSTRUMENTOS FINANCIEROS DE LA CIUDADANA {nombre}, CÉDULA {identificador}:"
    else:
        return f"INSTRUMENTOS FINANCIEROS DEL CIUDADANO {nombre}, CÉDULA {identificador}:"

def aplicar_formato_tabla(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        r'<w:tblBorders %s>'
        r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'  <w:left w:val="none"/>'
        r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'  <w:right w:val="none"/>'
        r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E6E6E6"/>'
        r'  <w:insideV w:val="none"/>'
        r'</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def sombrear_celda(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


# --- LÓGICA MÓDULO 1 ---
def procesar_modulo1(uploaded_file):
    doc_origen = Document(uploaded_file)
    lineas = []
    for p in doc_origen.paragraphs:
        if p.text.strip(): lineas.append(p.text.strip())
            
    for table in doc_origen.tables:
        for row in table.rows:
            textos_celda = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if textos_celda: lineas.append(",".join(textos_celda))

    titular_nombre, titular_id = "ALBERT JOSÉ ZAMBRANO ARELLANO", "V18760494"
    for l in lineas:
        m_tit = re.search(r'Persona Natural:\s*\n*([^,]+),([VEJG]\d+)', l, re.IGNORECASE)
        if m_tit:
            titular_nombre, titular_id = m_tit.group(1).strip(), m_tit.group(2).strip()

    cuentas_datos, recibidas_raw, enviadas_raw = [], [], []
    consolidado_por_anio = {
        "2025": {"cant_r": 0, "monto_r": 0.0, "cant_e": 0, "monto_e": 0.0},
        "2026": {"cant_r": 0, "monto_r": 0.0, "cant_e": 0, "monto_e": 0.0}
    }

    for l in lineas:
        match_cuenta = re.search(r'\b(\d{20})\b', l)
        if match_cuenta:
            nro_cuenta = match_cuenta.group(1)
            cuentas_datos.append([obtener_banco_por_cuenta(nro_cuenta), nro_cuenta, "TITULAR"])
            continue

        if "TRANSFERENCIA RECIBIDA" in l.upper() or "TRANSFERENCIA ENVIADA" in l.upper():
            match_rif = re.search(r'\b([VEJG]\d{5,10})\b', l, re.IGNORECASE)
            if match_rif:
                rif_encontrado = match_rif.group(1).upper()
                match_monto = re.search(r'\"([\d\.,]+)\"|,\s*([\d\.,]+)\s*,', l)
                monto_str = match_monto.group(1) if (match_monto and match_monto.group(1)) else (match_monto.group(2) if match_monto else "")
                monto = limpiar_monto_venezuela(monto_str)
                
                anio_detectado = "2026"
                match_fecha = re.search(r'\b\d{1,2}/\d{1,2}/(2025|2026)\b', l)
                if match_fecha: anio_detectado = match_fecha.group(1)
                
                texto_previo = l.split(rif_encontrado)[0].strip().rstrip(',')
                elementos_previos = texto_previo.split(',')
                nombre_origen = elementos_previos[-1].strip().replace('"', '') if elementos_previos else "SINFIRMA"
                if not nombre_origen or re.search(r'\d{2,4}', nombre_origen): nombre_origen = "SINFIRMA"

                registro = {"nombre": nombre_origen, "id": rif_encontrado, "cant": 1, "monto": monto}

                if "RECIBIDA" in l.upper():
                    recibidas_raw.append(registro)
                    if anio_detectado in consolidado_por_anio:
                        consolidado_por_anio[anio_detectado]["cant_r"] += 1
                        consolidado_por_anio[anio_detectado]["monto_r"] += monto
                else:
                    enviadas_raw.append(registro)
                    if anio_detectado in consolidado_por_anio:
                        consolidado_por_anio[anio_detectado]["cant_e"] += 1
                        consolidado_por_anio[anio_detectado]["monto_e"] += monto

    def consolidar_m1(lista_raw):
        agrupados = {}
        for item in lista_raw:
            clave = item["id"]
            if clave in agrupados:
                agrupados[clave]["cant"] += 1
                agrupados[clave]["monto"] += item["monto"]
            else:
                agrupados[clave] = item.copy()
        t_monto = sum(i["monto"] for i in agrupados.values())
        t_cant = sum(i["cant"] for i in agrupados.values())
        salida = []
        for i in agrupados.values():
            pct = (i["monto"] / t_monto * 100) if t_monto > 0 else 0.0
            salida.append([i["nombre"], i["id"], str(i["cant"]), formatear_monto_venezuela(i["monto"]), f"{pct:.2f}%".replace('.', ',')])
        salida.sort(key=lambda x: limpiar_monto_venezuela(x[3]), reverse=True)
        return salida, t_cant, t_monto

    r_tabla, c_tot_r, m_tot_r = consolidar_m1(recibidas_raw)
    e_tabla, c_tot_e, m_tot_e = consolidar_m1(enviadas_raw)

    doc = Document()
    _configurar_documento(doc)
    doc.add_paragraph().add_run(formatear_titulo_persona(titular_nombre, titular_id)).bold = True

    if cuentas_datos:
        vistas, unicas = set(), []
        for c in cuentas_datos:
            if c[1] not in vistas:
                vistas.add(c[1])
                unicas.append(c)
        _crear_tabla_cuentas(doc, unicas)

    _crear_tabla_consolidada_m1(doc, consolidado_por_anio, c_tot_r, m_tot_r, c_tot_e, m_tot_e)

    if r_tabla or e_tabla:
        doc.add_paragraph().add_run("TRANSFERENCIAS ELECTRÓNICAS (DETALLADAS):").bold = True
        if r_tabla: _crear_tabla_detallada(doc, "TRANSFERENCIAS RECIBIDAS", r_tabla, c_tot_r, m_tot_r)
        if e_tabla: _crear_tabla_detallada(doc, "TRANSFERENCIAS ENVIADAS", e_tabla, c_tot_e, m_tot_e)

    return doc, titular_id


# --- LÓGICA MÓDULO 2 ---
def procesar_modulo2(uploaded_file):
    doc_origen = Document(uploaded_file)
    titular_nombre, titular_id = "PERSONA EVALUADA", "DOCUMENTO"
    
    for p in doc_origen.paragraphs:
        texto_p = p.text.upper()
        if "NOMBRE" in texto_p and ("IDENTIFIC" in texto_p or "CEDULA" in texto_p or "C.I" in texto_p):
            partes = [parte.strip() for parte in p.text.split(',')]
            for parte in partes:
                if parte.upper().startswith("NOMBRE") and len(parte.split())>1:
                    titular_nombre = parte.split(maxsplit=1)[1].strip().upper()
                elif "IDENTIFIC" in parte.upper() or "CEDULA" in parte.upper() or "C.I" in parte.upper():
                    if len(parte.split())>1:
                        titular_id = parte.split(maxsplit=1)[1].strip().upper().replace('.', '').replace('-', '').replace(' ', '')
            break
            
    cuentas_datos, consolidado_datos, enviadas_raw, recibidas_raw = [], [], [], []
    orden_transferencias = 1 

    for table in doc_origen.tables:
        if not table.rows: continue
        encabezados = [cell.text.strip().upper().replace('\n', ' ') for cell in table.rows[0].cells]
        encabezados_unidos = " ".join(encabezados)

        if "NRO. CUENTA" in encabezados_unidos or "FIRMANTES" in encabezados_unidos:
            idx_cta = next((i for i, h in enumerate(encabezados) if "CUENTA" in h), -1)
            idx_firm = next((i for i, h in enumerate(encabezados) if "FIRMANTE" in h or "TITULAR" in h), -1)
            for row in table.rows[1:]:
                celdas = [c.text.strip() for c in row.cells]
                if len(celdas) > idx_cta and idx_cta != -1:
                    cuenta = re.sub(r'\D', '', celdas[idx_cta])
                    firm = celdas[idx_firm].strip() if (idx_firm != -1 and idx_firm < len(celdas)) else "TITULAR"
                    if len(cuenta) == 20: cuentas_datos.append([obtener_banco_por_cuenta(cuenta), cuenta, firm or "TITULAR"])
            continue

        if "TIPO" in encabezados_unidos and "CANTIDAD" in encabezados_unidos and "MONTO" in encabezados_unidos:
            for row in table.rows[1:]:
                celdas = [c.text.strip() for c in row.cells]
                if len(celdas) >= 3 and celdas[0]: consolidado_datos.append(celdas[:3])
            continue

        if "TIPO DE CLIENTE" in encabezados_unidos and "IDENTIFICACIÓN" in encabezados_unidos:
            idx_nom = next((i for i, h in enumerate(encabezados) if "NOMBRE" in h or "RAZÓN" in h), -1)
            idx_tip = next((i for i, h in enumerate(encabezados) if "TIPO" in h), -1)
            idx_ide = next((i for i, h in enumerate(encabezados) if "IDENTIFIC" in h), -1)
            idx_mon = next((i for i, h in enumerate(encabezados) if "MONTO" in h), -1)
            for row in table.rows[1:]:
                celdas = [c.text.strip() for c in row.cells]
                if len(celdas) > max(idx_nom, idx_tip, idx_ide, idx_mon) and "TOTAL" not in celdas[0].upper():
                    nombre, tipo_cl, identif, m_str = celdas[idx_nom].strip(), celdas[idx_tip].upper().strip(), celdas[idx_ide].strip(), celdas[idx_mon].strip()
                    if identif and tipo_cl:
                        ci_rif = f"{tipo_cl}{identif}".replace('-', '').replace('.', '').replace(' ', '')
                        reg = {"nombre": nombre, "id": ci_rif, "cant": 1, "monto": limpiar_monto_venezuela(m_str)}
                        if orden_transferencias == 1: enviadas_raw.append(reg)
                        else: recibidas_raw.append(reg)
            orden_transferencias += 1 
            continue

    def consolidar_m2(lista_raw):
        agrupados = {}
        for item in lista_raw:
            clave = item["id"]
            if clave in agrupados:
                agrupados[clave]["cant"] += 1
                agrupados[clave]["monto"] += item["monto"]
                if not agrupados[clave]["nombre"] and item["nombre"]: agrupados[clave]["nombre"] = item["nombre"]
            else: agrupados[clave] = item.copy()
        t_monto = sum(i["monto"] for i in agrupados.values())
        t_cant = sum(i["cant"] for i in agrupados.values())
        salida = []
        for i in agrupados.values():
            pct = (i["monto"] / t_monto * 100) if t_monto > 0 else 0.0
            salida.append([i["nombre"], i["id"], str(i["cant"]), formatear_monto_venezuela(i["monto"]), f"{pct:.2f}%".replace('.', ',')])
        salida.sort(key=lambda x: limpiar_monto_venezuela(x[3]), reverse=True)
        return salida, t_cant, t_monto

    doc = Document()
    _configurar_documento(doc)
    doc.add_paragraph().add_run(formatear_titulo_persona(titular_nombre, titular_id)).bold = True

    if cuentas_datos: _crear_tabla_cuentas(doc, cuentas_datos)
    
    if consolidado_datos:
        doc.add_paragraph().add_run("TRANSFERENCIAS ELECTRÓNICAS (CONSOLIDADAS):").bold = True
        t2 = doc.add_table(rows=1, cols=3)
        aplicar_formato_tabla(t2)
        for idx, text in enumerate(["TIPO", "CANTIDAD", "MONTO"]):
            t2.cell(0, idx).text = text
            t2.cell(0, idx).paragraphs[0].runs[0].font.bold = True
            sombrear_celda(t2.cell(0, idx), "F2F2F2")
        for fila in consolidado_datos:
            r_dat = t2.add_row().cells
            r_dat[0].text, r_dat[1].text, r_dat[2].text = fila[0], fila[1], fila[2]
            if "TOTAL" in fila[0].upper():
                for cell in r_dat: cell.paragraphs[0].runs[0].font.bold = True
        doc.add_paragraph()

    if enviadas_raw or recibidas_raw:
        doc.add_paragraph().add_run("TRANSFERENCIAS ELECTRÓNICAS (DETALLADAS):").bold = True
        if enviadas_raw:
            e_tabla, c_tot_e, m_tot_e = consolidar_m2(enviadas_raw)
            _crear_tabla_detallada(doc, "TRANSFERENCIAS ENVIADAS", e_tabla, c_tot_e, m_tot_e)
        if recibidas_raw:
            r_tabla, c_tot_r, m_tot_r = consolidar_m2(recibidas_raw)
            _crear_tabla_detallada(doc, "TRANSFERENCIAS RECIBIDAS", r_tabla, c_tot_r, m_tot_r)

    return doc, titular_id


# --- LÓGICA MÓDULO 3 ---
def procesar_modulo3(uploaded_file):
    def extraer_anio(celdas, anio_contexto):
        for c in celdas:
            text = str(c).strip()
            m_fecha_4d = re.search(r'\b\d{1,2}[/-]\d{1,2}[/-](202[4-9])\b|\b(202[4-9])[/-]\d{1,2}[/-]\d{1,2}\b', text)
            if m_fecha_4d: return int(m_fecha_4d.group(1) or m_fecha_4d.group(2))
            m_fecha_2d = re.search(r'\b\d{1,2}[/-]\d{1,2}[/-](2[4-9])\b', text)
            if m_fecha_2d: return 2000 + int(m_fecha_2d.group(1))
            m_anio_4d = re.search(r'\b(202[4-9])\b', text)
            if m_anio_4d: return int(m_anio_4d.group(1))
        return anio_contexto

    def consolidar_m3(lista_raw):
        agrupados = {}
        for item in lista_raw:
            clave = item["id"]
            if clave in agrupados:
                agrupados[clave]["cant"] += item["cant"]
                agrupados[clave]["monto"] += item["monto"]
                if not agrupados[clave]["nombre"] and item["nombre"]: agrupados[clave]["nombre"] = item["nombre"]
            else: agrupados[clave] = item.copy()
        t_monto = sum(i["monto"] for i in agrupados.values())
        t_cant = sum(i["cant"] for i in agrupados.values())
        salida = []
        for i in agrupados.values():
            pct = (i["monto"] / t_monto * 100) if t_monto > 0 else 0.0
            salida.append([i["nombre"], i["id"], str(i["cant"]), formatear_monto_venezuela(i["monto"]), f"{pct:.2f}%".replace('.', ',')])
        salida.sort(key=lambda x: limpiar_monto_venezuela(x[3]), reverse=True)
        return salida, t_cant, t_monto

    doc_origen = Document(uploaded_file)
    titular_nombre, titular_id = "PERSONA EVALUADA", "DOCUMENTO"
    cuentas_datos, enviadas_raw, recibidas_raw = [], [], []
    orden_transferencias = 1 
    
    for table in doc_origen.tables:
        for i, row in enumerate(table.rows):
            texto_linea = " ".join([c.text.strip().upper() for c in row.cells])
            if "NOMBRES Y APELLIDOS" in texto_linea or "PERSONA EVALUADA" in texto_linea or "NOMBRE O RAZÓN SOCIAL" in texto_linea:
                data_row = table.rows[i+1] if i+1 < len(table.rows) else row
                celdas_data = [c.text.strip().upper() for c in data_row.cells]
                for c in celdas_data:
                    match_id = re.search(r'\b[VJEG]-?\d{6,10}\b', c)
                    if match_id: titular_id = match_id.group(0).replace('-', '')
                if len(celdas_data) >= 3 and "NOMBRE" not in celdas_data[1]:
                    titular_nombre = celdas_data[1]
                    if titular_id == "DOCUMENTO": titular_id = celdas_data[2].replace('-', '')
                break
        if titular_nombre != "PERSONA EVALUADA": break

    for table in doc_origen.tables:
        if len(table.rows) < 2: continue
        head_text = " ".join([c.text.strip().upper() for r in table.rows[:2] for c in r.cells])
        anios_en_tabla = set()
        for r in table.rows:
            for c in r.cells:
                t_celda = c.text.strip()
                for a in re.findall(r'\b(202[4-9])\b', t_celda): anios_en_tabla.add(int(a))
                for a in re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-](2[4-9])\b', t_celda): anios_en_tabla.add(2000 + int(a))
        
        anio_contexto = 2025
        match_h_4d = re.search(r'\b(202[4-9])\b', head_text)
        if match_h_4d: anio_contexto = int(match_h_4d.group(1))
        elif anios_en_tabla: anio_contexto = list(anios_en_tabla)[0]

        if "NRO. CUENTA" in head_text or "CUENTA" in head_text:
            idx_firm = next((i for i, h in enumerate(table.rows[0].cells) if "FIRMA" in h.text.upper() or "TITULAR" in h.text.upper()), -1)
            for row in table.rows[1:]:
                celdas = [c.text.strip() for c in row.cells]
                cuenta = next((re.sub(r'\D', '', c) for c in celdas if len(re.sub(r'\D', '', c)) == 20), None)
                if cuenta:
                    firmante = celdas[idx_firm].strip().upper() if (idx_firm != -1 and idx_firm < len(celdas) and celdas[idx_firm]) else "TITULAR"
                    cuentas_datos.append([obtener_banco_por_cuenta(cuenta), cuenta, firmante])
            continue

        if ("TRANSFERENCIAS RECIBIDAS" in head_text and "TRANSFERENCIAS ENVIADAS" in head_text) or ("AÑO" in head_text and "RECIBIDAS" in head_text):
            for row in table.rows:
                celdas = [c.text.strip() for c in row.cells]
                if len(celdas) >= 5:
                    nombre = celdas[0].strip().upper()
                    if not nombre or "NOMBRE" in nombre or "TOTAL" in nombre or "AÑO" in nombre or "CANTIDAD" in nombre: continue
                    if nombre.isdigit() and len(nombre) == 4:
                        anio_fila, c_rec, m_rec, c_env, m_env = int(nombre), int(celdas[1]) if celdas[1].isdigit() else 0, limpiar_monto_venezuela(celdas[2]), int(celdas[3]) if celdas[3].isdigit() else 0, limpiar_monto_venezuela(celdas[4])
                        if m_rec > 0 or c_rec > 0: recibidas_raw.append({"nombre": "CONSOLIDADO", "id": "S/I", "cant": c_rec, "monto": m_rec, "anio": anio_fila})
                        if m_env > 0 or c_env > 0: enviadas_raw.append({"nombre": "CONSOLIDADO", "id": "S/I", "cant": c_env, "monto": m_env, "anio": anio_fila})
                        continue
                    identif = celdas[1].strip().upper().replace('-', '').replace('.', '').replace(' ', '')
                    if not re.search(r'\d', identif): continue
                    m_rec, m_env = limpiar_monto_venezuela(celdas[3]), limpiar_monto_venezuela(celdas[5]) if len(celdas) > 5 else 0.0
                    c_rec, c_env = int(celdas[2]) if celdas[2].isdigit() else 1, int(celdas[4]) if len(celdas) > 4 and celdas[4].isdigit() else 1
                    anio_fila = extraer_anio(celdas, anio_contexto)
                    if m_rec > 0: recibidas_raw.append({"nombre": nombre, "id": identif, "cant": c_rec, "monto": m_rec, "anio": anio_fila})
                    if m_env > 0: enviadas_raw.append({"nombre": nombre, "id": identif, "cant": c_env, "monto": m_env, "anio": anio_fila})
            continue

        if "TIPO DE CLIENTE" in head_text and "IDENTIFICACIÓN" in head_text:
            idx_nom = next((i for i, h in enumerate(table.rows[0].cells) if "NOMBRE" in h.text.upper() or "RAZÓN" in h.text.upper()), -1)
            idx_tip = next((i for i, h in enumerate(table.rows[0].cells) if "TIPO" in h.text.upper()), -1)
            idx_ide = next((i for i, h in enumerate(table.rows[0].cells) if "IDENTIFIC" in h.text.upper()), -1)
            idx_mon = next((i for i, h in enumerate(table.rows[0].cells) if "MONTO" in h.text.upper()), -1)
            is_recibidas = "RECIBIDA" in head_text or orden_transferencias == 2
            
            for row in table.rows[1:]:
                celdas = [c.text.strip() for c in row.cells]
                if len(celdas) > max(idx_nom, idx_tip, idx_ide, idx_mon) and idx_nom != -1 and idx_mon != -1:
                    nombre = celdas[idx_nom].strip().upper()
                    if "TOTAL" in nombre or not nombre: continue
                    tipo = celdas[idx_tip].strip().upper() if idx_tip != -1 else ""
                    ide = celdas[idx_ide].strip().upper().replace('-', '').replace('.', '').replace(' ', '')
                    monto, anio_fila = limpiar_monto_venezuela(celdas[idx_mon]), extraer_anio(celdas, anio_contexto)
                    if monto > 0:
                        reg = {"nombre": nombre, "id": f"{tipo}{ide}" if tipo else ide, "cant": 1, "monto": monto, "anio": anio_fila}
                        if is_recibidas: recibidas_raw.append(reg)
                        else: enviadas_raw.append(reg)
            orden_transferencias = 2 
            continue

    datos_consolidado = {}
    for r in recibidas_raw:
        anio = r.get("anio") or 2025
        if anio not in datos_consolidado: datos_consolidado[anio] = {'rec_c': 0, 'rec_m': 0.0, 'env_c': 0, 'env_m': 0.0}
        datos_consolidado[anio]['rec_c'] += r['cant']
        datos_consolidado[anio]['rec_m'] += r['monto']
    for e in enviadas_raw:
        anio = e.get("anio") or 2025
        if anio not in datos_consolidado: datos_consolidado[anio] = {'rec_c': 0, 'rec_m': 0.0, 'env_c': 0, 'env_m': 0.0}
        datos_consolidado[anio]['env_c'] += e['cant']
        datos_consolidado[anio]['env_m'] += e['monto']

    r_tabla, c_tot_r, m_tot_r = consolidar_m3([i for i in recibidas_raw if i["id"] != "S/I"])
    e_tabla, c_tot_e, m_tot_e = consolidar_m3([i for i in enviadas_raw if i["id"] != "S/I"])

    doc = Document()
    _configurar_documento(doc)
    doc.add_paragraph().add_run(formatear_titulo_persona(titular_nombre, titular_id)).bold = True

    if cuentas_datos:
        vistas, unicas = set(), []
        for f in cuentas_datos:
            if f[1] not in vistas:
                vistas.add(f[1])
                unicas.append(f)
        _crear_tabla_cuentas(doc, unicas)

    if datos_consolidado: _crear_tabla_consolidada_m3(doc, datos_consolidado)

    if r_tabla or e_tabla:
        doc.add_paragraph().add_run("TRANSFERENCIAS ELECTRÓNICAS (DETALLADAS):").bold = True
        if r_tabla: _crear_tabla_detallada(doc, "TRANSFERENCIAS RECIBIDAS", r_tabla, c_tot_r, m_tot_r)
        if e_tabla: _crear_tabla_detallada(doc, "TRANSFERENCIAS ENVIADAS", e_tabla, c_tot_e, m_tot_e)

    return doc, titular_id


# --- UTILIDADES DE DISEÑO DE WORD ---
def _configurar_documento(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = 'Arial', Pt(10)

def _crear_tabla_cuentas(doc, cuentas):
    t = doc.add_table(rows=1, cols=3)
    aplicar_formato_tabla(t)
    for idx, text in enumerate(['ENTIDADES BANCARIAS', 'NRO. CUENTA', 'TITULAR / FIRMANTE']):
        t.cell(0, idx).text = text
        t.cell(0, idx).paragraphs[0].runs[0].font.bold = True
        sombrear_celda(t.cell(0, idx), "F2F2F2")
    for f in cuentas:
        row = t.add_row().cells
        row[0].text, row[1].text, row[2].text = f[0], f[1], f[2]
    doc.add_paragraph()

def _crear_tabla_detallada(doc, titulo, lineas, c_t, m_t):
    doc.add_paragraph().add_run(titulo).bold = True
    t = doc.add_table(rows=1, cols=5)
    aplicar_formato_tabla(t)
    for idx, text in enumerate(["NOMBRE O RAZÓN SOCIAL", "C.I O RIF", "Nº TRANSF", "MONTO", "%"]):
        t.cell(0, idx).text = text
        t.cell(0, idx).paragraphs[0].runs[0].font.bold = True
        sombrear_celda(t.cell(0, idx), "F2F2F2")
    for f in lineas:
        row = t.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = f[0], f[1], f[2], f[3], f[4]
    r_f = t.add_row().cells
    r_f[0].text, r_f[2].text, r_f[3].text, r_f[4].text = "TOTAL GENERAL", str(c_t), formatear_monto_venezuela(m_t), "100,00%"
    for cell in r_f:
        if cell.text: cell.paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

def _crear_tabla_consolidada_m1(doc, por_anio, c_r, m_r, c_e, m_e):
    doc.add_paragraph().add_run("TRANSFERENCIAS ELECTRÓNICAS (CONSOLIDADAS) EN BOLÍVARES:").bold = True
    t2 = doc.add_table(rows=2, cols=5)
    aplicar_formato_tabla(t2)
    t2.cell(0, 0).text = "AÑO"
    t2.cell(0, 1).merge(t2.cell(0, 2)).text = "RECIBIDAS"
    t2.cell(0, 3).merge(t2.cell(0, 4)).text = "ENVIADAS"
    for idx, text in enumerate(["", "CANTIDAD", "MONTO BS", "CANTIDAD", "MONTO BS"]):
        if text: t2.cell(1, idx).text = text
    for r_idx in [0, 1]:
        for cell in t2.rows[r_idx].cells:
            if cell.text:
                cell.paragraphs[0].runs[0].font.bold = True
                sombrear_celda(cell, "F2F2F2")
    for anio in ["2025", "2026"]:
        d = por_anio[anio]
        r_dat = t2.add_row().cells
        r_dat[0].text, r_dat[1].text, r_dat[2].text, r_dat[3].text, r_dat[4].text = anio, str(d["cant_r"]), formatear_monto_venezuela(d["monto_r"]), str(d["cant_e"]), formatear_monto_venezuela(d["monto_e"])
    r_tot = t2.add_row().cells
    r_tot[0].text, r_tot[1].text, r_tot[2].text, r_tot[3].text, r_tot[4].text = "TOTAL GENERAL", str(c_r), formatear_monto_venezuela(m_r), str(c_e), formatear_monto_venezuela(m_e)
    for cell in r_tot: cell.paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

def _crear_tabla_consolidada_m3(doc, datos_años):
    doc.add_paragraph().add_run("TRANSFERENCIAS ELECTRÓNICAS: (CONSOLIDADO)").bold = True
    t = doc.add_table(rows=2, cols=5)
    aplicar_formato_tabla(t)
    for i, txt in enumerate(["AÑO", "RECIBIDAS", "", "ENVIADAS", ""]):
        if txt: 
            t.cell(0, i).text = txt
            if txt != "AÑO": t.cell(0, i).merge(t.cell(0, i+1))
            else: t.cell(0, 0).merge(t.cell(1, 0))
            t.cell(0, i).paragraphs[0].runs[0].font.bold = True
            t.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            sombrear_celda(t.cell(0, i), "F2F2F2")
    for i, txt in enumerate(["", "CANTIDAD", "MONTO (Bs.)", "CANTIDAD", "MONTO (Bs.)"]):
        if i > 0:
            t.cell(1, i).text = txt
            t.cell(1, i).paragraphs[0].runs[0].font.bold = True
            t.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            sombrear_celda(t.cell(1, i), "F2F2F2")
    tr_c, tr_m, te_c, te_m = 0, 0.0, 0, 0.0
    for anio in sorted(datos_años.keys()):
        d = datos_años[anio]
        fila = t.add_row().cells
        fila[0].text, fila[1].text, fila[2].text, fila[3].text, fila[4].text = str(anio), str(d['rec_c']), formatear_monto_venezuela(d['rec_m']), str(d['env_c']), formatear_monto_venezuela(d['env_m'])
        tr_c += d['rec_c']; tr_m += d['rec_m']; te_c += d['env_c']; te_m += d['env_m']
    r_total = t.add_row().cells
    r_total[0].paragraphs[0].add_run("TOTAL GENERAL").bold = True
    r_total[1].paragraphs[0].add_run(str(tr_c)).bold = True
    r_total[2].paragraphs[0].add_run(formatear_monto_venezuela(tr_m)).bold = True
    r_total[3].paragraphs[0].add_run(str(te_c)).bold = True
    r_total[4].paragraphs[0].add_run(formatear_monto_venezuela(te_m)).bold = True
    doc.add_paragraph()


# --- INTERFAZ DE USUARIO (STREAMLIT) ---
st.markdown("---")
st.subheader("Selecciona el Módulo de Procesamiento:")

# Creamos 3 columnas para poner las imágenes y las casillas una al lado de la otra
col1, col2, col3 = st.columns(3)

with col1:
    # Puedes cambiar esta URL por el enlace de tu propia imagen
    st.image("foto1.png", width=80) 
    mod1 = st.checkbox("Módulo 1: Detección Básica")

with col2:
    st.image("foto2.png", width=80) 
    mod2 = st.checkbox("Módulo 2: Automática Dinámica")

with col3:
    st.image("foto3.png", width=80) 
    mod3 = st.checkbox("Módulo 3: Integral por Años")

# Validamos que el usuario seleccione solo UNA casilla a la vez
casillas_seleccionadas = sum([mod1, mod2, mod3])

if casillas_seleccionadas > 1:
    st.warning("⚠️ Por favor, selecciona **solo un** módulo a la vez para evitar conflictos.")
    procesamiento_permitido = False
elif casillas_seleccionadas == 0:
    st.info("👆 Selecciona al menos un módulo arriba para continuar.")
    procesamiento_permitido = False
else:
    procesamiento_permitido = True

st.markdown("---")
archivo_subido = st.file_uploader("Sube tu archivo Word (.docx)", type=["docx"])

# Botón de procesamiento
if st.button("Procesar Archivo", type="primary"):
    if not procesamiento_permitido:
        st.error("Revisa la selección del módulo antes de procesar.")
    elif archivo_subido is None:
        st.warning("Por favor, sube un archivo Word antes de procesar.")
    else:
        try:
            with st.spinner("Analizando documento..."):
                # Ejecutamos la lógica según la casilla marcada
                if mod1:
                    doc_final, titular_id = procesar_modulo1(archivo_subido)
                    sufijo = "Basico"
                elif mod2:
                    doc_final, titular_id = procesar_modulo2(archivo_subido)
                    sufijo = "Dinamico"
                elif mod3:
                    doc_final, titular_id = procesar_modulo3(archivo_subido)
                    sufijo = "Completo"

                # Guardar en memoria para descarga
                buffer_salida = io.BytesIO()
                doc_final.save(buffer_salida)
                buffer_salida.seek(0)
                nombre_archivo_salida = f"Analisis_Financiero_{sufijo}_{titular_id}.docx"

            st.success("¡Documento procesado con éxito! 🎉")
            st.download_button(
                label="📥 Descargar Reporte Generado",
                data=buffer_salida,
                file_name=nombre_archivo_salida,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Ocurrió un error al procesar el archivo: {e}")
